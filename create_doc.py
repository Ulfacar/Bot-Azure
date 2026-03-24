from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Стили
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Заголовок
title = doc.add_heading('Подключение WhatsApp для бота Ton Azure', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# Вступление
p = doc.add_paragraph()
p.add_run('Бектур, здравствуйте!').bold = True
doc.add_paragraph(
    'Этот документ объясняет как мы подключаем WhatsApp к боту отеля, '
    'почему выбрали именно этот способ и что нужно от вас.'
)

# Раздел 1
doc.add_heading('Почему не через Meta (Facebook)?', level=2)
doc.add_paragraph(
    'Подключение WhatsApp напрямую через Meta требует:'
)
bullets = [
    'Верификацию бизнес-аккаунта Facebook — процесс занимает от нескольких дней до нескольких недель',
    'Подтверждённое бизнес-портфолио в Meta Business Manager',
    'Прохождение проверки документов компании',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'На данный момент аккаунт Facebook ещё новый и не прошёл все этапы верификации. '
    'Чтобы не терять время и запустить бота как можно скорее, мы используем альтернативный сервис.'
)

doc.add_paragraph(
    'В будущем, когда аккаунт Facebook будет полностью готов, мы сможем перейти на официальный '
    'API Meta — при необходимости.'
)

# Раздел 2
doc.add_heading('Что мы используем — Wappi.pro', level=2)
doc.add_paragraph(
    'Wappi.pro — это сервис для подключения WhatsApp к ботам и бизнес-системам. '
    'Бот будет работать точно так же: клиент пишет в WhatsApp — бот отвечает автоматически.'
)

doc.add_heading('Преимущества:', level=3)
benefits = [
    'Быстрое подключение — бот заработает в WhatsApp в тот же день',
    'Не нужна верификация Facebook',
    'Безлимитные сообщения — без ограничений на количество',
    'Стоимость — 700 рублей в месяц (~1200 сом)',
]
for b in benefits:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
p.add_run('Важно: ').bold = True
p.add_run(
    'так как сервис работает через обычный WhatsApp (не Business API), '
    'рекомендуется использовать отдельный номер телефона только для бота. '
    'Также советуем иметь 2 запасных сим-карты на случай если потребуется сменить номер.'
)

# Раздел 3
doc.add_heading('Что нужно от вас', level=2)

doc.add_heading('Шаг 1. Купить SIM-карту для бота', level=3)
doc.add_paragraph(
    'Нужна отдельная SIM-карта (новый номер) специально для WhatsApp бота. '
    'Это не должен быть чей-то личный номер. Также купите 2 запасных SIM-карты — на всякий случай для бизнеса.'
)

doc.add_heading('Шаг 2. Зарегистрировать WhatsApp на этом номере', level=3)
doc.add_paragraph(
    'Вставьте SIM-карту в любой телефон, установите WhatsApp и пройдите регистрацию '
    '(подтверждение по SMS). После этого WhatsApp должен быть активен на этом номере.'
)

doc.add_heading('Шаг 3. Зарегистрироваться на Wappi.pro', level=3)
steps = [
    'Зайдите на сайт wappi.pro',
    'Нажмите «Начать бесплатно» (есть 5 дней бесплатного периода)',
    'Создайте аккаунт (email + пароль)',
    'В личном кабинете добавьте профиль WhatsApp — отсканируйте QR-код с телефона где установлен WhatsApp бота',
    'Выберите тариф «Базовый» — 700 руб/мес',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Шаг 4. Отправить нам данные', level=3)
doc.add_paragraph('После регистрации и подключения отправьте нам:')
data_needed = [
    'API Token (находится в личном кабинете Wappi.pro → Настройки профиля)',
    'Profile ID (там же)',
    'Номер телефона который подключили',
]
for d in data_needed:
    doc.add_paragraph(d, style='List Bullet')

doc.add_paragraph(
    'Мы пропишем эти данные в систему и бот начнёт работать в WhatsApp.'
)

# Раздел 4
doc.add_heading('Итого расходы', level=2)

table = doc.add_table(rows=3, cols=2)
table.style = 'Light Shading Accent 1'
table.rows[0].cells[0].text = 'Что'
table.rows[0].cells[1].text = 'Стоимость'
table.rows[1].cells[0].text = 'Wappi.pro (ежемесячно)'
table.rows[1].cells[1].text = '700 руб/мес (~1200 сом)'
table.rows[2].cells[0].text = '3 SIM-карты (разово)'
table.rows[2].cells[1].text = '~300-500 сом'

doc.add_paragraph('')

# Раздел 5
doc.add_heading('Напоминание по другим вопросам', level=2)
reminders = [
    'Telegram ID менеджера — чтобы бот отправлял уведомления о новых клиентах. Пусть менеджер напишет боту @userinfobot в Telegram, он покажет ID. Отправьте нам этот ID.',
    'Пополнить OpenRouter на $5-10 — для подключения более умной AI модели (Gemini). Реквизиты для пополнения отправим отдельно.',
    'Single comfort — это отдельный тип номера или обычный двухместный для одного гостя? Нужно для корректных ответов бота.',
]
for r in reminders:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('С уважением, команда разработки ASystem').italic = True

# Сохранение
output_path = r'c:\Users\alanb\OneDrive\Рабочий стол\WhatsApp подключение Ton Azure.docx'
doc.save(output_path)
print(f'Done: {output_path}')
